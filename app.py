import streamlit as st
import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH # WD_SHADING removed to resolve import error
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement # For setting table cell borders
import io # To handle file in memory

# --- Helper function to set table cell borders ---
def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    Args:
        cell: The docx.table._Cell object.
        kwargs: Keyword arguments for borders (top, bottom, left, right).
                Each value should be a dict with 'sz' (size in eighths of a point),
                'color' (RGBColor object or hex string like '000000'), and 'val' (border style).
    """
    tcPr = cell._element.get_or_add_tcPr()

    for border_name in ['top', 'bottom', 'left', 'right']:
        if border_name in kwargs:
            border_props = kwargs[border_name]
            bdr = OxmlElement(f'w:{border_name}Bdr')

            border_val = border_props.get('val', 'single')
            bdr.set(qn('w:val'), border_val)
            bdr.set(qn('w:sz'), str(border_props.get('sz', 12))) # Default 1.5pt

            if border_val != 'nil':
                # Attempt to get color, defaulting to black RGBColor if not provided
                color_input = border_props.get('color', RGBColor(0, 0, 0))
                hex_color = '000000' # Default black hex as fallback

                if isinstance(color_input, RGBColor):
                    try:
                        # Attempt to get hex from RGBColor.rgb (standard way)
                        # The .rgb property returns a Bytes object, which is subscriptable
                        hex_color = f'{color_input.rgb[0]:02X}{color_input.rgb[1]:02X}{color_input.rgb[2]:02X}'
                    except AttributeError:
                        # If RGBColor object truly has no .rgb attribute (due to corrupted lib),
                        # fall back to default black hex.
                        pass # hex_color remains '000000'
                elif isinstance(color_input, str):
                    # If color was provided as a hex string directly
                    hex_color = color_input
                
                bdr.set(qn('w:color'), hex_color)
            
            tcPr.append(bdr)

# --- Function to generate DOCX ---
def generate_invoice_docx(invoice_data):
    document = Document()

    # Set document margins (typically 1.27 cm / 0.5 inches or 1.5cm / 0.6 inches for a clean look)
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # --- Header Section: "TAX INVOICE/DELIVERY NOTE" ---
    title = document.add_paragraph('TAX INVOICE/DELIVERY NOTE')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.runs[0]
    run_title.font.name = 'Arial' # Common clean font
    run_title.bold = True
    run_title.font.size = Pt(16) # Slightly reduced from 18, check template
    # Add spacing after title
    title.paragraph_format.space_after = Pt(12)

    # --- Table for Contact and Invoice Details ---
    # This table typically has no visible borders
    header_info_table = document.add_table(rows=6, cols=4)
    header_info_table.autofit = False
    header_info_table.allow_autofit = False
    
    # Precise column widths for header table (adjust these by small increments if needed)
    header_info_table.columns[0].width = Cm(3.0) # Labels like "To:", "Address:"
    header_info_table.columns[1].width = Cm(6.5) # Values like Company Name, Address
    header_info_table.columns[2].width = Cm(3.5) # Labels like "Date:", "SSS Invoice No:"
    header_info_table.columns[3].width = Cm(5.0) # Values like Date, Invoice No

    # Populate cells
    data_rows = [
        ["To:", invoice_data['to_company'], "Date:", invoice_data['invoice_date']],
        ["Address:", invoice_data['customer_address'], "SSS Invoice No:", invoice_data['sss_invoice_no']],
        ["Tel:", invoice_data['customer_tel'], "Customer VAT No.:", invoice_data['customer_vat_no']],
        ["ATTN:", invoice_data['attn_person'], "", ""],
        ["Email:", invoice_data['customer_email'], "", ""],
        ["Customer PO#:", invoice_data['customer_po'], "", ""]
    ]

    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = header_info_table.rows[r_idx].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9) # Standard text size for details
            # Ensure no borders
            set_cell_border(cell, top={'sz': 0, 'val': 'nil'}, bottom={'sz': 0, 'val': 'nil'},
                                left={'sz': 0, 'val': 'nil'}, right={'sz': 0, 'val': 'nil'})

    # Handle merges for the right-hand empty cells in ATTN, Email, PO# rows
    header_info_table.rows[3].cells[2].merge(header_info_table.rows[3].cells[3])
    header_info_table.rows[4].cells[2].merge(header_info_table.rows[4].cells[3])
    header_info_table.rows[5].cells[2].merge(header_info_table.rows[5].cells[3])

    # Add spacing after header info table
    document.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Line Items Table ---
    item_table = document.add_table(rows=1, cols=7)
    item_table.autofit = False
    item_table.allow_autofit = False

    # Set precise column widths for line items table
    item_table.columns[0].width = Cm(1.2)  # No
    item_table.columns[1].width = Cm(7.8)  # Description (adjusted)
    item_table.columns[2].width = Cm(2.2)  # Unit price (adjusted)
    item_table.columns[3].width = Cm(1.2)  # BHD (Unit) (adjusted)
    item_table.columns[4].width = Cm(1.2)  # Qty (adjusted)
    item_table.columns[5].width = Cm(2.2)  # Total price (adjusted)
    item_table.columns[6].width = Cm(1.2)  # BHD (Total) (adjusted)

    # Apply 'Table Grid' style for standard borders
    item_table.style = 'Table Grid'

    # Add header row
    item_table_headers = ["No", "Description", "Unit price", "BHD", "Qty", "Total price", "BHD"]
    hdr_cells = item_table.rows[0].cells
    for i, header_text in enumerate(item_table_headers):
        cell = hdr_cells[i]
        cell.text = header_text
        paragraph = cell.paragraphs[0]
        # Align headers
        if header_text in ["No", "Qty", "BHD"]:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif header_text in ["Unit price", "Total price"]:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Apply bold, font, and size for headers
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.bold = True
            run.font.size = Pt(9)
        # Removed shading application to resolve persistent import error
        # cell.shading.fill = WD_SHADING.TEXTURE_NONE
        # cell.shading.foreground = RGBColor(0xD9, 0xD9, 0xD9)

    # Add item rows
    for i, item in enumerate(invoice_data['line_items']):
        row_cells = item_table.add_row().cells
        # No.
        cell_no = row_cells[0]
        cell_no.text = str(i + 1)
        cell_no.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Description
        cell_desc = row_cells[1]
        cell_desc.text = item['description']
        cell_desc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Unit price
        cell_unit_price = row_cells[2]
        cell_unit_price.text = f"{item['unit_price']:.3f}"
        cell_unit_price.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # BHD (Unit)
        cell_bhd_unit = row_cells[3]
        cell_bhd_unit.text = "BHD"
        cell_bhd_unit.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Qty
        cell_qty = row_cells[4]
        cell_qty.text = str(item['quantity'])
        cell_qty.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Total price
        cell_total_price = row_cells[5]
        cell_total_price.text = f"{item['total_price']:.3f}"
        cell_total_price.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # BHD (Total)
        cell_bhd_total = row_cells[6]
        cell_bhd_total.text = "BHD"
        cell_bhd_total.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply font size for all cells in item rows
        for j in range(7):
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)

    # Add 5 empty rows if current items are less than 5 to maintain a consistent table height
    # This might be needed if your template always shows a minimum number of rows
    # For dynamic content, this can be skipped. Assuming template has fixed empty rows after items.
    num_existing_items = len(invoice_data['line_items'])
    min_rows_to_show = 5 # Or whatever your template's minimum is
    if num_existing_items < min_rows_to_show:
        for _ in range(min_rows_to_show - num_existing_items):
            row_cells = item_table.add_row().cells
            for cell in row_cells:
                # Ensure empty cells also have borders and consistent font if typed into
                set_cell_border(cell, top={'sz': 12}, bottom={'sz': 12},
                                left={'sz': 12}, right={'sz': 12})
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)


    # --- Summary Section (Subtotal, VAT, Grand Total) ---
    # This table is typically positioned to align its right edge with the line items table
    summary_table = document.add_table(rows=3, cols=2)
    summary_table.autofit = False
    summary_table.allow_autofit = False

    # Calculate required width for the first column to align right
    # Total width of line item table = sum of its column widths = 1.2+7.8+2.2+1.2+1.2+2.2+1.2 = 17cm
    # Width of Total Price + BHD columns = 2.2 + 1.2 = 3.4cm
    # So, first column of summary table should be (17 - 3.4) = 13.6cm to align
    summary_table.columns[0].width = Cm(13.6) # Label column (Subtotal:, VAT @ 10%:, Grand Total:)
    summary_table.columns[1].width = Cm(3.4)  # Value column (Amounts)

    # Ensure no borders for summary table cells, as per typical invoice design.
    for row in summary_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={'sz': 0, 'val': 'nil'}, bottom={'sz': 0, 'val': 'nil'},
                                left={'sz': 0, 'val': 'nil'}, right={'sz': 0, 'val': 'nil'})

    # Row 1: Subtotal
    subtotal_cell_label = summary_table.rows[0].cells[0]
    subtotal_cell_label.text = "Subtotal:"
    subtotal_cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtotal_cell_label.paragraphs[0].paragraph_format.space_before = Pt(6) # Add small space before
    subtotal_cell_label.paragraphs[0].paragraph_format.space_after = Pt(6) # Add small space after

    subtotal_cell_value = summary_table.rows[0].cells[1]
    subtotal_cell_value.text = f"{invoice_data['subtotal']:.3f} BHD"
    subtotal_cell_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtotal_cell_value.paragraphs[0].paragraph_format.space_before = Pt(6)
    subtotal_cell_value.paragraphs[0].paragraph_format.space_after = Pt(6)

    # Row 2: VAT
    vat_cell_label = summary_table.rows[1].cells[0]
    vat_cell_label.text = "VAT @ 10%:"
    vat_cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    vat_cell_label.paragraphs[0].paragraph_format.space_before = Pt(6)
    vat_cell_label.paragraphs[0].paragraph_format.space_after = Pt(6)

    vat_cell_value = summary_table.rows[1].cells[1]
    vat_cell_value.text = f"{invoice_data['vat_amount']:.3f} BHD"
    vat_cell_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    vat_cell_value.paragraphs[0].paragraph_format.space_before = Pt(6)
    vat_cell_value.paragraphs[0].paragraph_format.space_after = Pt(6)

    # Row 3: Grand Total
    grand_total_cell_label = summary_table.rows[2].cells[0]
    grand_total_cell_label.text = "Grand Total in BHD"
    grand_total_cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    grand_total_cell_label.paragraphs[0].paragraph_format.space_before = Pt(12) # More space before grand total
    grand_total_cell_label.paragraphs[0].paragraph_format.space_after = Pt(12) # More space after grand total


    grand_total_cell_value = summary_table.rows[2].cells[1]
    grand_total_cell_value.text = f"{invoice_data['grand_total']:.3f} BHD"
    grand_total_cell_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    grand_total_cell_value.paragraphs[0].paragraph_format.space_before = Pt(12)
    grand_total_cell_value.paragraphs[0].paragraph_format.space_after = Pt(12)

    # Apply bold and larger font size for Grand Total
    for p in grand_total_cell_label.paragraphs:
        for run in p.runs:
            run.font.name = 'Arial'
            run.bold = True
            run.font.size = Pt(12) # Larger font

    for p in grand_total_cell_value.paragraphs:
        for run in p.runs:
            run.font.name = 'Arial'
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0) # Black color

    document.add_paragraph().paragraph_format.space_after = Pt(20) # More space after summary

    # --- Payment Details (Static Text) ---
    payment_details_para = document.add_paragraph()
    payment_details_para.add_run("Payment will be remitted to the following account:").font.name = 'Arial'
    payment_details_para.add_run("\nSalahuddin Softtech Solutions, Arab Bank, A/C No.: 2002-146072-510,").font.name = 'Arial'
    payment_details_para.add_run("\nIBAN: BH31 ARAB 0200 2146072510, Swift Code: ARAB BHBM,").font.name = 'Arial'
    payment_details_para.add_run("\nAddress: Arab Bank Plc. , P.O Box: 395, Manama, Kingdom of Bahrain.").font.name = 'Arial'
    for run in payment_details_para.runs:
        run.font.size = Pt(9)
    payment_details_para.paragraph_format.space_after = Pt(20)

    # --- Terms and Conditions ---
    terms_heading = document.add_paragraph('Terms and Conditions')
    terms_heading.runs[0].bold = True
    terms_heading.runs[0].font.size = Pt(11) # Slightly smaller heading
    terms_heading.runs[0].font.name = 'Arial'
    terms_heading.paragraph_format.space_after = Pt(6) # Small space after heading

    terms_para1 = document.add_paragraph()
    terms_para1.add_run("Payment terms: ").bold = True
    terms_para1.add_run(invoice_data['payment_terms'])
    for run in terms_para1.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
    terms_para1.paragraph_format.space_after = Pt(6)

    terms_para2 = document.add_paragraph("Title and property at all above remain ours until full payment is received and we reserve the rights to withdraw goods/services if not paid for when due.")
    for run in terms_para2.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
    terms_para2.paragraph_format.space_after = Pt(6)


    terms_para3 = document.add_paragraph("Signing this document implies acceptance of these terms.")
    for run in terms_para3.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9)
    terms_para3.paragraph_format.space_after = Pt(20) # More space after terms

    # --- Receipt/Acknowledgement ---
    document.add_paragraph("Received by").paragraph_format.space_after = Pt(40) # Add significant space after 'Received by'

    # Create a 3-column table for Name, Date, Signature alignment
    ack_signature_table = document.add_table(rows=1, cols=3)
    ack_signature_table.autofit = False
    ack_signature_table.allow_autofit = False

    # Set widths to spread them out appropriately across the page
    # Total content width approx 18cm. Divide equally.
    col_width_ack = Cm(18.0 / 3) # Roughly 6cm per column
    ack_signature_table.columns[0].width = col_width_ack
    ack_signature_table.columns[1].width = col_width_ack
    ack_signature_table.columns[2].width = col_width_ack

    # Remove all borders for this table
    for row in ack_signature_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={'sz': 0, 'val': 'nil'}, bottom={'sz': 0, 'val': 'nil'},
                                left={'sz': 0, 'val': 'nil'}, right={'sz': 0, 'val': 'nil'})

    ack_cells = ack_signature_table.rows[0].cells
    
    # Name column
    ack_name_para = ack_cells[0].paragraphs[0]
    ack_name_para.add_run("(Name)").font.name = 'Arial'
    ack_name_para.add_run("\n").add_underline = True # Add underline for the line
    ack_name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ack_name_para.runs:
        run.font.size = Pt(9)

    # Date column
    ack_date_para = ack_cells[1].paragraphs[0]
    ack_date_para.add_run("(Date)").font.name = 'Arial'
    ack_date_para.add_run("\n").add_underline = True
    ack_date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ack_date_para.runs:
        run.font.size = Pt(9)

    # Signature column
    ack_sig_para = ack_cells[2].paragraphs[0]
    ack_sig_para.add_run("(Signature)").font.name = 'Arial'
    ack_sig_para.add_run("\n").add_underline = True
    ack_sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ack_sig_para.runs:
        run.font.size = Pt(9)

    # Add a blank paragraph for vertical spacing after the signature lines
    document.add_paragraph().paragraph_format.space_after = Pt(20)


    # Final company details
    doc_para_company = document.add_paragraph("For SALAHUDDIN SOFTTECH SOLUTIONS")
    doc_para_name = document.add_paragraph("Jobin George")
    doc_para_title = document.add_paragraph("Operations Manager")

    # Apply font formatting to these final paragraphs
    for para in [doc_para_company, doc_para_name, doc_para_title]:
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10) # Standard size for final contact
        # Reduce space after for these lines to be closer
        para.paragraph_format.space_after = Pt(3) # Small space between lines


    # Save the document to an in-memory byte stream
    byte_io = io.BytesIO()
    document.save(byte_io)
    byte_io.seek(0)
    return byte_io

# --- Streamlit Form Layout ---
st.title("Invoice Generator")

st.subheader("Line Items")

# Initialize line_items in session_state if not present
if 'line_items' not in st.session_state:
    st.session_state.line_items = []
if 'invoice_data_ready' not in st.session_state:
    st.session_state.invoice_data_ready = False # Flag to control preview visibility

# Buttons to add/remove items (OUTSIDE the form)
if st.button("Add New Item"):
    st.session_state.line_items.append({"description": "", "unit_price": 0.000, "quantity": 1, "total_price": 0.000})
    st.session_state.invoice_data_ready = False # Hide preview if items change
    st.session_state.generated_docx_data = None # Clear generated docx

total_subtotal = 0.000
items_to_remove = []

for i, item in enumerate(st.session_state.line_items):
    st.write(f"--- Item {i+1} ---")
    item_cols = st.columns([4, 2, 1, 1]) # Description, Unit Price, Qty, Remove Button
    with item_cols[0]:
        item['description'] = st.text_input("Description", item['description'], key=f"desc_{i}")
    with item_cols[1]:
        item['unit_price'] = st.number_input("Unit price (BHD)", min_value=0.000, value=float(item['unit_price']), format="%.3f", key=f"price_{i}")
    with item_cols[2]:
        item['quantity'] = st.number_input("Qty", min_value=1, value=int(item['quantity']), step=1, key=f"qty_{i}")

    item['total_price'] = item['unit_price'] * item['quantity']
    total_subtotal += item['total_price']

    st.text(f"Total Price for Item {i+1}: {item['total_price']:.3f} BHD")
    if item_cols[3].button(f"Remove Item {i+1}", key=f"remove_{i}"):
        items_to_remove.append(i)
        st.session_state.invoice_data_ready = False # Hide preview if items change
        st.session_state.generated_docx_data = None # Clear generated docx


for i in sorted(items_to_remove, reverse=True):
    del st.session_state.line_items[i]
if items_to_remove:
    st.experimental_rerun()


# --- Main Invoice Form (for header details and final submission) ---
with st.form("invoice_form"):
    st.subheader("Customer & Invoice Details")

    col1_form, col2_form = st.columns(2)
    with col1_form:
        to_company = st.text_input("To (Company Name)", "ABC Company W.L.L.")
        customer_address = st.text_area("Address", "Building 123, Road 456, Block 789, Manama, Bahrain")
        customer_tel = st.text_input("Tel", "+973 17XXXXXX")
        attn_person = st.text_input("ATTN", "Mr. John Doe")
        customer_email = st.text_input("Email", "john.doe@abccompany.com")
        customer_po = st.text_input("Customer PO#", "PO-12345")

    with col2_form:
        invoice_date = st.date_input("Date", datetime.date.today())
        current_date_str = datetime.date.today().strftime("%y%m%d")
        sss_invoice_no = st.text_input("SSS Invoice No", f"SSS-{current_date_str}-001")
        customer_vat_no = st.text_input("Customer VAT No.", "VAT123456789")

    vat_rate = 0.10
    vat_amount = total_subtotal * vat_rate
    grand_total = total_subtotal + vat_amount

    st.subheader("Summary")
    st.markdown(f"**Subtotal:** {total_subtotal:.3f} BHD")
    st.markdown(f"**VAT @ 10%:** {vat_amount:.3f} BHD")
    st.markdown(f"## **Grand Total:** {grand_total:.3f} BHD")


    st.subheader("Terms and Conditions")
    payment_terms = st.text_area("Payment terms", "30 days from invoice date")

    # --- Submit Button for the Form ---
    submitted = st.form_submit_button("Generate Invoice Preview") # Renamed button

    if submitted:
        if not st.session_state.line_items:
            st.warning("Please add at least one line item to generate the invoice.")
            # Clear existing generated data if any
            st.session_state.invoice_data_ready = False
            st.session_state.generated_docx_data = None
            st.session_state.generated_docx_filename = None
        else:
            invoice_data = {
                "to_company": to_company,
                "customer_address": customer_address,
                "customer_tel": customer_tel,
                "attn_person": attn_person,
                "customer_email": customer_email,
                "customer_po": customer_po,
                "invoice_date": invoice_date.strftime("%d-%m-%Y"),
                "sss_invoice_no": sss_invoice_no,
                "customer_vat_no": customer_vat_no,
                "line_items": st.session_state.line_items,
                "subtotal": total_subtotal,
                "vat_amount": vat_amount,
                "grand_total": grand_total,
                "payment_terms": payment_terms
            }

            st.success("Invoice data collected successfully! Displaying preview...")

            # Store invoice_data in session_state to be accessible for preview and download
            st.session_state.invoice_data_for_preview = invoice_data
            st.session_state.invoice_data_ready = True # Flag to show preview section
            st.session_state.generated_docx_data = None # Clear previously generated docx on new preview generation

# --- Invoice Preview Section (appears if invoice_data is ready) ---
if 'invoice_data_ready' in st.session_state and st.session_state.invoice_data_ready:
    st.markdown("---")
    st.subheader("Invoice Preview")
    preview_data = st.session_state.invoice_data_for_preview

    st.markdown(f"### **TAX INVOICE/DELIVERY NOTE**")
    st.markdown("---")
    st.markdown(f"**To:** {preview_data['to_company']} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; **Date:** {preview_data['invoice_date']}")
    st.markdown(f"**Address:** {preview_data['customer_address']} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; **SSS Invoice No:** {preview_data['sss_invoice_no']}")
    st.markdown(f"**Tel:** {preview_data['customer_tel']} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; **Customer VAT No.:** {preview_data['customer_vat_no']}")
    st.markdown(f"**ATTN:** {preview_data['attn_person']}")
    st.markdown(f"**Email:** {preview_data['customer_email']}")
    st.markdown(f"**Customer PO#:** {preview_data['customer_po']}")
    st.markdown("---")

    st.markdown("### Line Items:")
    # Create a simple table for line items in markdown
    line_item_table_md = "| No | Description | Unit Price | Qty | Total Price |\n|---|---|---|---|---|\n"
    for i, item in enumerate(preview_data['line_items']):
        line_item_table_md += f"| {i+1} | {item['description']} | {item['unit_price']:.3f} BHD | {item['quantity']} | {item['total_price']:.3f} BHD |\n"
    st.markdown(line_item_table_md)
    st.markdown("---")

    st.markdown(f"**Subtotal:** {preview_data['subtotal']:.3f} BHD")
    st.markdown(f"**VAT @ 10%:** {preview_data['vat_amount']:.3f} BHD")
    st.markdown(f"## **Grand Total:** {preview_data['grand_total']:.3f} BHD")
    st.markdown("---")

    st.markdown("### Terms and Conditions")
    st.markdown(f"**Payment terms:** {preview_data['payment_terms']}")
    st.markdown("Title and property at all above remain ours until full payment is received and we reserve the rights to withdraw goods/services if not paid for when due.")
    st.markdown("Signing this document implies acceptance of these terms.")
    st.markdown("---")
    st.markdown("Received by")
    # Using non-breaking spaces for basic alignment in markdown
    st.markdown("\n\n(Name) &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; (Date) &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; (Signature)")
    st.markdown("---")
    st.markdown("For SALAHUDDIN SOFTTECH SOLUTIONS")
    st.markdown("Jobin George")
    st.markdown("Operations Manager")

    st.write("---")
    st.subheader("Generate & Download Document")

    # Only generate DOCX when download button is clicked to avoid re-generating on every rerun
    if st.button("Generate & Download Word Document"):
        docx_file_bytes = generate_invoice_docx(preview_data) # Use the data from session_state

        st.session_state.generated_docx_data = docx_file_bytes
        st.session_state.generated_docx_filename = f"Invoice_{preview_data['sss_invoice_no']}.docx" # Use the SSS Invoice No from preview data

        # This download button must be defined AFTER its data is set,
        # and only appears after the "Generate & Download" button is clicked.
        st.download_button(
            label="Click here to Download Invoice (Word)", # Changed label to indicate it's ready
            data=st.session_state.generated_docx_data,
            file_name=st.session_state.generated_docx_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="final_download_button" # Unique key for this download button
        )
        st.success("Word document generated and ready for download!")
        st.info("PDF generation requires external libraries/tools for exact formatting. We can explore options for this next if the Word document is satisfactory.")
# --- Download button (remains here as well for initial setup, but activated by logic above) ---
# This part of the download button logic is removed because the download button is now conditional
# inside the "Generate & Download Word Document" block.
# if 'generated_docx_data' in st.session_state and st.session_state.generated_docx_data is not None:
#     st.download_button(
#         label="Download Invoice (Word)",
#         data=st.session_state.generated_docx_data,
#         file_name=st.session_state.generated_docx_filename,
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#     )
